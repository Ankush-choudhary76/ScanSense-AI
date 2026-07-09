from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Optional
import base64
import os
import io
import PyPDF2
from PIL import Image

# Initialize Database correctly from backend directory
from .database.db_manager import (
    init_db, create_session, save_message, get_chat_history,
    get_all_sessions, delete_session, update_session_context,
    get_session_context, save_memory_summary, get_recent_memory_summaries
)
from .utils.ai_assistant import stream_chat_response, is_medical_content, summarize_session

app = FastAPI(title="ScanSense AI API")

# Add CORS middleware for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins since frontend is served from same server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
def startup_event():
    try:
        # Make sure we're initializing the SQLite tables
        init_db()
        print("✅ Database initialized successfully")
    except Exception as e:
        print(f"❌ Database initialization error: {str(e)}")
        raise

# Models
class ChatRequest(BaseModel):
    session_id: Optional[int] = None
    message: str
    context: Optional[str] = None
    content_type: Optional[str] = None

class SessionCreate(BaseModel):
    title: Optional[str] = None

@app.post("/api/sessions")  # Create new chat session
def new_session(session: SessionCreate):
    session_id = create_session(session.title)
    return {"session_id": session_id}

@app.get("/api/sessions")   # List all chat sessions
def list_sessions():
    return {"sessions": get_all_sessions()}

@app.get("/api/sessions/{session_id}/history")  # Get chat history for a session 
def get_history(session_id: int):
    return {"history": get_chat_history(session_id)}

@app.delete("/api/sessions/{session_id}")   # Delete a chat session
def delete_chat_session(session_id: int):
    delete_session(session_id)
    return {"status": "deleted"}

@app.post("/api/upload")    # Upload a image or PDF/Word/Image file and extract content
async def upload_file(file: UploadFile = File(...)):
    """ Endpoint to extract content from files and return context """
    try:
        content = await file.read()
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        extracted_context = ""
        content_type = ""
        encoded_image = None
        
        if file_ext == '.pdf':
            content_type = 'pdf'
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                extracted_context += page.extract_text()
                
        elif file_ext == '.docx':
            content_type = 'pdf'  # Treat document text context same as pdf
            import docx
            doc = docx.Document(io.BytesIO(content))
            text_runs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_runs.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_runs.append(" | ".join(row_text))
            extracted_context = "\n".join(text_runs)
            
        elif file_ext == '.doc':
            content_type = 'pdf'  # Treat document text context same as pdf
            import re
            # Extract plain text sequences from old binary .doc files as a fallback
            printable_sequences = re.findall(rb'[a-zA-Z0-9\s\.,;:!\?\-\(\)\'\"\r\n]{4,}', content)
            extracted_context = "\n".join([seq.decode('utf-8', errors='ignore').strip() for seq in printable_sequences])
                
        elif file_ext in ['.png', '.jpg', '.jpeg', '.jpe', '.webp', '.bmp', '.gif', '.tiff']:
            content_type = 'image'
            image = Image.open(io.BytesIO(content))
            buffered = io.BytesIO()
            image.convert('RGB').save(buffered, format="JPEG")
            encoded_image = base64.b64encode(buffered.getvalue()).decode('utf-8')
            extracted_context = encoded_image
            
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {file_ext}")

        return {
            "status": "success",
            "content_type": content_type,
            "context": extracted_context,
            "filename": file.filename
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat")  # send a message and get AI response (with SSE streaming)
async def chat_endpoint(request: ChatRequest, background_tasks: BackgroundTasks):
    if not request.session_id:
        request.session_id = create_session(request.message[:30] + "...")
    
    # Context Logic
    if request.context and request.content_type:
        update_session_context(request.session_id, request.context, request.content_type)
        current_context = request.context
        current_content_type = request.content_type
    else:
        stored = get_session_context(request.session_id)
        current_context = stored.get("context")
        current_content_type = stored.get("content_type")

    # Save the user's message to SQLite
    save_message(request.session_id, "user", request.message)
    
    # Grab short-term history (current session)
    chat_history = get_chat_history(request.session_id)

    # Fetch long-term memory (summaries of past sessions)
    long_term_memories = get_recent_memory_summaries(
        exclude_session_id=request.session_id, limit=5
    )

    # Generate response (SSE streaming)
    def sse_event_generator():
        full_response = ""
        yield f"data: SESSION_ID:{request.session_id}\n\n"
        
        for chunk in stream_chat_response(
            chat_history=chat_history,
            context=current_context,
            content_type=current_content_type,
            user_text=request.message,
            long_term_memories=long_term_memories
        ):
            full_response += chunk
            # Standard SSE has issues with raw newlines in 'data:' chunks.
            # We replace \n with \r (carriage return) which frontend will swap back.
            safe_chunk = chunk.replace("\n", "\r")
            yield f"data: {safe_chunk}\n\n"
        
        # Save assistant reply
        save_message(request.session_id, "assistant", full_response)

        # Background: generate & save session summary for long-term memory
        updated_history = get_chat_history(request.session_id)
        if len(updated_history) >= 2:
            summary = summarize_session(updated_history)
            if summary:
                save_memory_summary(request.session_id, summary)

        yield "data: [DONE]\n\n"

    return StreamingResponse(sse_event_generator(), media_type="text/event-stream")

# --- Static File Serving (Connect Frontend) ---

# Get the absolute path to the frontend directory
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIST = os.path.join(BASE_DIR, "..", "frontend")

# Mount the static files (compiled React app or vanilla HTML)
if os.path.exists(FRONTEND_DIST):
    # We no longer strictly need /assets mounted if everything is in frontend root, 
    # but we can mount the whole frontend root safely as static exceptions or keep it 
    # handled by the catchall. Let's just rely on the catch-all for simplicity 
    # or mount specific dirs if needed.
    # We will just remove the specific assets mount since it's not needed for vanilla js structure,
    # and let the catch-all resolve everything.

    @app.get("/{rest_of_path:path}")
    async def serve_frontend(rest_of_path: str):
        # Serve the file if it exists (like style.css, app.js, images in assets etc.)
        file_path = os.path.join(FRONTEND_DIST, rest_of_path)
        if os.path.isfile(file_path):
            return FileResponse(file_path)
        # Otherwise, serve index.html
        return FileResponse(os.path.join(FRONTEND_DIST, "index.html"))
else:
    print(f"Warning: Frontend build directory not found at {FRONTEND_DIST}. Run 'npm run build' in the frontend folder.")